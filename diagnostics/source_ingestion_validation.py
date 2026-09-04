import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from fabric_rag.config import (
    DATABASE_PATH,
    DOCUMENTS_DIRECTORY,
    EXPECTED_EMBEDDING_DIMENSIONS,
)
from fabric_rag.documents import load_documents, process_documents
from fabric_rag.knowledge_base import rebuild_database
from fabric_rag.management import (
    get_knowledge_base_management_status,
    indexed_chunk_identity,
    remove_source_document,
    save_source_document,
)


def raises(expected_exception, operation):
    try:
        operation()
    except expected_exception:
        return True
    return False


def pdf_bytes(text):
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    page = writer.add_blank_page(width=612, height=792)

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_reference = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {NameObject("/F1"): font_reference}
            )
        }
    )

    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    stream = DecodedStreamObject()
    stream.set_data(
        f"BT /F1 12 Tf 72 720 Td ({escaped_text}) Tj ET".encode("latin-1")
    )
    page[NameObject("/Contents")] = writer._add_object(stream)

    output = BytesIO()
    writer.write(output)
    return output.getvalue()


def docx_bytes():
    document = Document()
    document.add_heading("Architecture Guide", level=1)
    document.add_paragraph("DOCX paragraphs remain in their original order.")
    document.add_paragraph("First implementation standard", style="List Bullet")
    document.add_paragraph("Second implementation standard", style="List Bullet")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def embedded_chunks_for(source_directory):
    processed = process_documents(load_documents(source_directory))
    chunks = [chunk for document in processed for chunk in document["chunks"]]
    return [
        {
            **chunk,
            "embedding": [float(index + 1)]
            + [0.0] * (EXPECTED_EMBEDDING_DIMENSIONS - 1),
        }
        for index, chunk in enumerate(chunks)
    ]


def rebuild_fixture(source_directory, database_path):
    rebuild_database(
        embedded_chunks_for(source_directory),
        database_path=database_path,
    )


def document_map(source_directory):
    return {
        document["source"]: document["raw_text"]
        for document in load_documents(source_directory)
    }


def current_corpus_identity():
    documents = load_documents(DOCUMENTS_DIRECTORY)
    processed = process_documents(documents)
    chunks = [chunk for document in processed for chunk in document["chunks"]]
    identity = sorted(
        (chunk["source"], chunk["chunk_number"], chunk["text"])
        for chunk in chunks
    )
    return documents, chunks, identity


def main():
    results = {}

    with tempfile.TemporaryDirectory(prefix="fabric-rag-ingestion-") as temp_root:
        root = Path(temp_root)
        sources = root / "sources"
        database = root / "rag.db"

        save_source_document(
            "utf8.md",
            "# UTF-8 Markdown\n\nA standard Markdown source.".encode("utf-8"),
            sources,
        )
        cp1254_text = "# İç Politika\n\nÇalışma belgeleri güvenli biçimde işlenir."
        save_source_document(
            "windows-1254.md",
            cp1254_text.encode("cp1254"),
            sources,
        )
        utf16_text = "UTF-16 source text with a deterministic byte-order mark."
        save_source_document(
            "utf16.txt",
            utf16_text.encode("utf-16"),
            sources,
        )
        save_source_document(
            "notes.txt",
            b"Plain text enters the shared chunking pipeline.",
            sources,
        )
        save_source_document(
            "guide.pdf",
            pdf_bytes("Text-based PDF content is extractable."),
            sources,
        )
        save_source_document("standards.docx", docx_bytes(), sources)

        loaded = document_map(sources)
        results["A. UTF-8 Markdown"] = "standard Markdown" in loaded["utf8.md"]
        results["B. Windows-1254 Markdown/text"] = (
            "İç Politika" in loaded["windows-1254.md"]
            and "Çalışma" in loaded["windows-1254.md"]
        )
        results["C. UTF-16 text"] = utf16_text in loaded["utf16.txt"]
        results["D. Plain-text ingestion"] = (
            "shared chunking pipeline" in loaded["notes.txt"]
        )
        results["E. Text-based PDF ingestion"] = (
            "Text-based PDF content is extractable" in loaded["guide.pdf"]
        )
        results["F. DOCX ingestion"] = (
            loaded["standards.docx"].startswith("# Architecture Guide")
            and "- First implementation standard" in loaded["standards.docx"]
            and "- Second implementation standard" in loaded["standards.docx"]
        )
        results["G. Unsupported extension rejected"] = raises(
            ValueError,
            lambda: save_source_document(
                "unsupported.html",
                b"<p>Unsupported</p>",
                sources,
            ),
        )
        results["H. Malformed PDF rejected"] = (
            raises(
                ValueError,
                lambda: save_source_document(
                    "broken.pdf",
                    b"not a pdf",
                    sources,
                ),
            )
            and not (sources / "broken.pdf").exists()
        )
        results["I. Malformed DOCX rejected"] = (
            raises(
                ValueError,
                lambda: save_source_document(
                    "broken.docx",
                    b"not a docx package",
                    sources,
                ),
            )
            and not (sources / "broken.docx").exists()
        )
        results["J. Same-name duplicate rejected"] = raises(
            FileExistsError,
            lambda: save_source_document(
                "notes.txt",
                b"Replacement text must not overwrite.",
                sources,
            ),
        )

        rebuild_fixture(sources, database)
        unchanged_status = get_knowledge_base_management_status(sources, database)
        results["L. Unchanged mixed corpus is up to date"] = not unchanged_status[
            "rebuild_required"
        ]

        save_source_document("added.txt", b"A newly added source.", sources)
        added_is_stale = get_knowledge_base_management_status(
            sources,
            database,
        )["rebuild_required"]
        rebuild_fixture(sources, database)

        (sources / "guide.pdf").write_bytes(
            pdf_bytes("The PDF changed while retaining the same filename.")
        )
        modified_is_stale = get_knowledge_base_management_status(
            sources,
            database,
        )["rebuild_required"]
        rebuild_fixture(sources, database)

        remove_source_document("standards.docx", sources)
        removed_is_stale = get_knowledge_base_management_status(
            sources,
            database,
        )["rebuild_required"]
        results["K. Mixed add/remove/modify requires rebuild"] = (
            added_is_stale and modified_is_stale and removed_is_stale
        )

    current_documents, current_chunks, current_identity = current_corpus_identity()
    results["M. Bundled Markdown corpus remains 8/90"] = (
        len(current_documents) == 8
        and len(current_chunks) == 90
        and all(document["source"].endswith(".md") for document in current_documents)
        and current_identity == indexed_chunk_identity(DATABASE_PATH)
    )

    print("MULTI-FORMAT SOURCE INGESTION VALIDATION")
    print("=" * 48)
    for name, passed in results.items():
        print(f"{name}: {'PASS' if passed else 'FAIL'}")

    passed_count = sum(results.values())
    print(f"\nPassed: {passed_count} / {len(results)}")
    if passed_count != len(results):
        raise RuntimeError("At least one source-ingestion validation failed.")


if __name__ == "__main__":
    main()
